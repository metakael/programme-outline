"""
API Routes for Programme Outline Generator (Continued)
----------------------------------------
Flask routes for handling API requests.
"""

# API Routes for Generated Outlines (continued)
@app.route('/api/generate', methods=['POST'])
def generate_outline():
    """Generate a new outline based on specifications"""
    # Get request data
    data = request.json
    
    if not data:
        return jsonify({
            'success': False,
            'error': 'No data provided'
        }), 400
    
    # Extract specifications
    title = data.get('title', 'New Workshop')
    objectives = data.get('objectives', '')
    total_duration = data.get('totalDuration', 120)
    segments = data.get('segments', [])
    style_adherence = float(data.get('styleAdherence', 0.8))
    reference_ids = data.get('referenceIds', [])
    
    # Create specifications dictionary
    specifications = {
        'title': title,
        'objectives': objectives,
        'total_duration': total_duration,
        'segments': segments
    }
    
    # Get reference outlines
    reference_outlines = []
    if reference_ids:
        reference_outlines = ReferenceOutline.query.filter(ReferenceOutline.id.in_(reference_ids)).all()
    else:
        # Use embeddings to find similar outlines if no specific references provided
        # This is a simplification - in a real app, we'd do proper embedding similarity search
        reference_outlines = ReferenceOutline.query.limit(3).all()
    
    # Generate the outline
    generated_content = outline_generator.generate_outline(
        specifications, reference_outlines, style_adherence
    )
    
    # Save the generated outline
    outline = GeneratedOutline(
        title=title,
        objectives=objectives,
        total_duration=total_duration,
        specifications=specifications,
        content=generated_content,
        reference_id=reference_outlines[0].id if reference_outlines else None,
        user_id=None  # Set user_id if authentication is implemented
    )
    
    db.session.add(outline)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'outline': outline.to_dict(),
        'content': generated_content
    })

@app.route('/api/outlines', methods=['GET'])
def get_outlines():
    """Get all generated outlines"""
    outlines = GeneratedOutline.query.all()
    return jsonify({
        'success': True,
        'outlines': [outline.to_dict() for outline in outlines]
    })

@app.route('/api/outlines/<int:outline_id>', methods=['GET'])
def get_outline(outline_id):
    """Get a specific generated outline"""
    outline = GeneratedOutline.query.get_or_404(outline_id)
    return jsonify({
        'success': True,
        'outline': outline.to_dict(),
        'content': outline.content
    })

@app.route('/api/outlines/<int:outline_id>', methods=['PUT'])
def update_outline(outline_id):
    """Update a generated outline"""
    outline = GeneratedOutline.query.get_or_404(outline_id)
    data = request.json
    
    if not data:
        return jsonify({
            'success': False,
            'error': 'No data provided'
        }), 400
    
    # Update fields
    if 'title' in data:
        outline.title = data['title']
    if 'content' in data:
        outline.content = data['content']
    if 'objectives' in data:
        outline.objectives = data['objectives']
    if 'totalDuration' in data:
        outline.total_duration = data['totalDuration']
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'outline': outline.to_dict(),
        'content': outline.content
    })

@app.route('/api/outlines/<int:outline_id>', methods=['DELETE'])
def delete_outline(outline_id):
    """Delete a generated outline"""
    outline = GeneratedOutline.query.get_or_404(outline_id)
    db.session.delete(outline)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'message': f'Outline {outline_id} deleted'
    })

@app.route('/api/outlines/<int:outline_id>/regenerate-segment', methods=['POST'])
def regenerate_segment(outline_id):
    """Regenerate a specific segment of an outline"""
    outline = GeneratedOutline.query.get_or_404(outline_id)
    data = request.json
    
    if not data:
        return jsonify({
            'success': False,
            'error': 'No data provided'
        }), 400
    
    segment_index = data.get('segmentIndex')
    segment_specs = data.get('segmentSpecs', {})
    
    if segment_index is None:
        return jsonify({
            'success': False,
            'error': 'No segment index provided'
        }), 400
    
    # Get reference outlines
    reference_ids = data.get('referenceIds', [])
    reference_outlines = []
    
    if reference_ids:
        reference_outlines = ReferenceOutline.query.filter(ReferenceOutline.id.in_(reference_ids)).all()
    elif outline.reference_id:
        reference_outlines = [ReferenceOutline.query.get(outline.reference_id)]
    else:
        reference_outlines = ReferenceOutline.query.limit(3).all()
    
    # Regenerate the segment
    updated_content = outline_generator.regenerate_segment(
        outline.content, segment_index, {'segment': segment_specs}, reference_outlines
    )
    
    # Update the outline
    outline.content = updated_content
    db.session.commit()
    
    return jsonify({
        'success': True,
        'outline': outline.to_dict(),
        'content': updated_content
    })

# Export Routes
@app.route('/api/outlines/<int:outline_id>/export', methods=['GET'])
def export_outline(outline_id):
    """Export an outline to the specified format"""
    outline = GeneratedOutline.query.get_or_404(outline_id)
    format_type = request.args.get('format', 'pdf')
    
    if format_type == 'pdf':
        return export_as_pdf(outline)
    elif format_type == 'docx':
        return export_as_docx(outline)
    elif format_type == 'text':
        return export_as_text(outline)
    else:
        return jsonify({
            'success': False,
            'error': f'Unsupported format: {format_type}'
        }), 400

def export_as_pdf(outline):
    """Export outline as PDF"""
    from fpdf import FPDF
    
    # Create PDF
    pdf = FPDF()
    pdf.add_page()
    
    # Add title
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, outline.title, 0, 1, 'C')
    
    # Add objectives
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, 'Objectives:', 0, 1, 'L')
    pdf.set_font('Arial', '', 12)
    
    # Split objectives into multiple lines if needed
    objectives = outline.objectives.strip().split('\n')
    for obj in objectives:
        pdf.multi_cell(0, 10, obj)
    
    pdf.ln(10)
    
    # Add content
    pdf.set_font('Arial', '', 12)
    
    # Replace bullet points and handle formatting
    content_lines = outline.content.strip().split('\n')
    for line in content_lines:
        line = line.strip()
        if not line:
            pdf.ln(5)
            continue
        
        # Check if line is a heading
        if line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            pdf.set_font('Arial', 'B', 12)
            pdf.multi_cell(0, 10, line)
            pdf.set_font('Arial', '', 12)
        # Check if line is a bullet point
        elif line.startswith('•'):
            pdf.set_x(15)  # Indent
            pdf.multi_cell(0, 10, line)
        else:
            pdf.multi_cell(0, 10, line)
    
    # Output PDF to a BytesIO object
    output = io.BytesIO()
    pdf.output(output, 'F')
    output.seek(0)
    
    # Create response
    return send_file(
        output,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f"{outline.title.replace(' ', '_')}.pdf"
    )

def export_as_docx(outline):
    """Export outline as DOCX"""
    from docx import Document
    
    # Create DOCX
    doc = Document()
    
    # Add title
    doc.add_heading(outline.title, level=1)
    
    # Add objectives
    doc.add_heading('Objectives:', level=2)
    doc.add_paragraph(outline.objectives)
    
    # Add content
    content_lines = outline.content.strip().split('\n')
    for line in content_lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if line is a heading
        if line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            doc.add_heading(line, level=2)
        # Check if line is a bullet point
        elif line.startswith('•'):
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(line[1:].strip())
        else:
            doc.add_paragraph(line)
    
    # Output DOCX to a BytesIO object
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Create response
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"{outline.title.replace(' ', '_')}.docx"
    )

def export_as_text(outline):
    """Export outline as plain text"""
    output = io.BytesIO()
    output.write(outline.content.encode('utf-8'))
    output.seek(0)
    
    # Create response
    return send_file(
        output,
        mimetype='text/plain',
        as_attachment=True,
        download_name=f"{outline.title.replace(' ', '_')}.txt"
    )

"""
Integration of PDF processing with API routes for Programme Outline Generator
"""

from app.pdf_improvements import enhance_pdf_extraction, extract_outline_structure

# Update this function in api_routes.py in the create_reference route
def process_uploaded_file(filepath, filename):
    """
    Process an uploaded file based on its type
    
    Args:
        filepath: Path to the uploaded file
        filename: Name of the uploaded file
        
    Returns:
        Extracted content as text
    """
    content = ""
    
    if filename.endswith('.docx'):
        from docx import Document
        doc = Document(filepath)
        content = "\n".join([para.text for para in doc.paragraphs])
        
    elif filename.endswith('.pdf'):
        # Use enhanced PDF extraction
        content = enhance_pdf_extraction(filepath)
        
    else:
        # For txt and md files
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
            
    return content

# Update this function in api_routes.py in the create_reference route
def parse_reference_outline(content, filepath=None, filename=None):
    """
    Parse a reference outline into structured data
    
    Args:
        content: Text content of the outline
        filepath: Optional path to the original file
        filename: Optional name of the file
        
    Returns:
        Parsed structure dictionary
    """
    # For PDF files, use the specialized extraction
    if filename and filename.endswith('.pdf') and filepath:
        structure = extract_outline_structure(content)
    else:
        # Use the standard parser for other file types
        from app.outline_parser import OutlineParser
        parser = OutlineParser()
        structure = parser.parse_outline(content)
    
    return structure

# Replace the related code in the create_reference route:
"""
@app.route('/api/references', methods=['POST'])
def create_reference():
    ...
    
    # Extract content based on file type
    content = process_uploaded_file(filepath, filename)
    
    # Parse the outline
    structure = parse_reference_outline(content, filepath, filename)
    
    # Generate embedding
    embedding = embedding_generator.generate_embedding(content, structure)
    
    # Create new reference outline
    title = structure.get('title', filename)
    reference = ReferenceOutline(
        title=title,
        description=request.form.get('description', ''),
        content=content,
        structure=structure,
        embedding=embedding
    )
    
    ...
"""
